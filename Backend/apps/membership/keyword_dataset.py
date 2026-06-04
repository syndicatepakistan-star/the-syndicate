"""Parse admin-uploaded keyword datasets: CSV, Word (.docx), or PDF.

Structured tables/lines are parsed locally. Unstructured PDF/DOCX/plain text is passed to OpenAI
to extract {category, keyword} seeds automatically when structured parsing yields no rows.
"""
from __future__ import annotations

import csv
import io
import sys
import uuid
import zipfile
from typing import Any

MIN_CHARS_FOR_AI_EXTRACTION = 80
MAX_CHARS_SENT_TO_AI = 18_000
# PDF loose parse can create thousands of comma-lines; prefer AI seeds above this.
MAX_STRUCTURED_ROWS_BEFORE_AI = 250
MAX_DATASET_ROWS = 500

VALID_CATEGORIES = frozenset({"business", "money", "power", "grooming", "others"})
_OLE_DOC_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


class KeywordDatasetParseError(ValueError):
    """Raised when a file cannot be interpreted as a keyword dataset."""


def normalize_category(raw: str) -> str:
    c = (raw or "").strip().lower()
    if c in VALID_CATEGORIES:
        return c
    aliases = {
        "finance": "money",
        "wealth": "money",
        "style": "grooming",
        "self-care": "grooming",
        "misc": "others",
        "other": "others",
        "general": "others",
    }
    return aliases.get(c, "others")


def _append_row(rows: list[dict[str, str]], row: dict[str, str | Any]) -> None:
    if not isinstance(row, dict):
        return
    norm = {(k or "").strip().lower(): (v or "").strip() for k, v in row.items()}
    cat = (
        norm.get("category")
        or norm.get("topic")
        or norm.get("pillar")
        or norm.get("bucket")
        or ""
    )
    kw = norm.get("keyword") or norm.get("keywords") or norm.get("phrase") or norm.get("term") or ""
    if not kw:
        vals = [str(x).strip() for x in row.values() if str(x).strip()]
        if len(vals) >= 2:
            cat, kw = vals[0], vals[1]
        elif vals:
            kw = vals[0]
    if not kw:
        return
    entry: dict[str, str] = {"category": normalize_category(cat), "keyword": kw[:500]}
    lvl = norm.get("level") or norm.get("tier") or norm.get("difficulty") or ""
    if lvl:
        entry["level"] = lvl[:24]
    title = norm.get("title") or norm.get("headline") or norm.get("name") or ""
    if title:
        entry["title"] = title[:500]
    desc = (
        norm.get("description")
        or norm.get("desc")
        or norm.get("summary")
        or norm.get("excerpt")
        or norm.get("subtitle")
        or norm.get("blurb")
        or ""
    )
    if desc:
        entry["description"] = desc[:900]
    source = (
        norm.get("source_text")
        or norm.get("source")
        or norm.get("content")
        or norm.get("body")
        or norm.get("text")
        or norm.get("article")
        or norm.get("details")
        or norm.get("notes")
        or norm.get("passage")
        or ""
    )
    if source:
        entry["source_text"] = source[:8000]
    rows.append(entry)


def _normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _parse_keyword_csv_loose(text: str) -> list[dict[str, str]]:
    """
    Fallback for exports with odd quoting / stray newlines: one data row per line,
    first comma separates category from keyword (keyword may contain commas if rare — prefer fixing CSV).
    """
    rows: list[dict[str, str]] = []
    for line in _normalize_newlines(text).split("\n"):
        line = line.strip()
        if not line:
            continue
        low = line.lower()
        if low.startswith("category") and "keyword" in low:
            continue
        if "," not in line:
            continue
        cat, kw = line.split(",", 1)
        cat, kw = cat.strip(), kw.strip().strip('"').strip()
        if kw:
            rows.append({"category": normalize_category(cat), "keyword": kw[:500]})
    return rows


def _decode_bytes_for_csv(raw: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("latin-1")


def _is_docx_zip(raw: bytes) -> bool:
    if not raw.startswith(b"PK"):
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            return "word/document.xml" in zf.namelist()
    except zipfile.BadZipFile:
        return False


def _is_xlsx_zip(raw: bytes) -> bool:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            names = zf.namelist()
            return "xl/workbook.xml" in names or any(n.startswith("xl/worksheets/") for n in names)
    except zipfile.BadZipFile:
        return False


def _extract_pdf_text(raw: bytes) -> str:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=raw, filetype="pdf")
        try:
            parts = [doc.load_page(i).get_text() for i in range(doc.page_count)]
        finally:
            doc.close()
        return "\n".join(parts)
    except Exception:
        pass
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw), strict=False)
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception as exc:
        raise KeywordDatasetParseError(f"Could not read this PDF: {exc}") from exc


def _parse_docx_bytes(raw: bytes) -> list[dict[str, str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise KeywordDatasetParseError("python-docx is required for Word files.") from exc

    rows: list[dict[str, str]] = []
    try:
        doc = Document(io.BytesIO(raw))
    except Exception as exc:
        raise KeywordDatasetParseError(f"Could not read this Word file: {exc}") from exc

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                _append_row(rows, {"category": cells[0], "keyword": cells[1]})
            elif len(cells) == 1 and "," in cells[0]:
                a, b = cells[0].split(",", 1)
                _append_row(rows, {"category": a, "keyword": b})

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        low = line.lower()
        if low.startswith("category") and "keyword" in low:
            continue
        if "," in line:
            a, b = line.split(",", 1)
            _append_row(rows, {"category": a.strip(), "keyword": b.strip().strip('"')})

    return rows


def _docx_plain_text(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise KeywordDatasetParseError("python-docx is required for Word files.") from exc
    try:
        doc = Document(io.BytesIO(raw))
    except Exception as exc:
        raise KeywordDatasetParseError(f"Could not read this Word file: {exc}") from exc
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells if (c.text or "").strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_keyword_csv_bytes(raw: bytes) -> list[dict[str, str]]:
    """
    Parse keyword CSV / plain text. Uses decoded text + newline='' pattern for csv.DictReader.
    """
    if not raw or not raw.strip():
        return []

    text = _decode_bytes_for_csv(raw)
    rows: list[dict[str, str]] = []
    dict_reader_kw: dict[str, Any] = {}
    if sys.version_info >= (3, 12):
        dict_reader_kw["strict"] = False

    stream = io.StringIO(text, newline="")
    try:
        reader = csv.DictReader(stream, **dict_reader_kw)
        if not reader.fieldnames:
            return []
        for row in reader:
            _append_row(rows, row)
        return rows
    except csv.Error:
        loose = _parse_keyword_csv_loose(text)
        return loose if loose else rows
    finally:
        stream.close()


def _try_structured_keyword_rows(raw: bytes, filename: str = "") -> list[dict[str, str]]:
    """
    Fast path: CSV / line-based category,keyword / docx tables only. Returns [] if nothing matched.
    """
    if not raw or not raw.strip():
        return []

    name = (filename or "").lower()

    if raw.startswith(_OLE_DOC_MAGIC):
        raise KeywordDatasetParseError(
            "Legacy Word .doc files are not supported. In Word use File → Save As → "
            "Word Document (.docx) or CSV UTF-8, then upload again."
        )

    if raw.startswith(b"%PDF") or name.endswith(".pdf"):
        text = _extract_pdf_text(raw)
        return _parse_keyword_csv_loose(text)

    if zipfile.is_zipfile(io.BytesIO(raw)):
        if _is_xlsx_zip(raw) and not _is_docx_zip(raw):
            raise KeywordDatasetParseError(
                "Excel .xlsx workbooks are not supported for this upload. "
                "Export as CSV UTF-8, or copy the table into Word (.docx) and upload that."
            )
        if _is_docx_zip(raw) or name.endswith(".docx"):
            return _parse_docx_bytes(raw)
        raise KeywordDatasetParseError(
            "This file is a ZIP archive, not a supported format. Upload CSV, Word (.docx), or PDF."
        )

    return parse_keyword_csv_bytes(raw)


def _extract_full_document_text(raw: bytes, filename: str = "") -> str:
    """Plain text for AI keyword extraction (PDF, DOCX, CSV/TXT body)."""
    if not raw or not raw.strip():
        return ""

    name = (filename or "").lower()

    if raw.startswith(_OLE_DOC_MAGIC):
        raise KeywordDatasetParseError(
            "Legacy Word .doc files are not supported. Save as .docx or CSV UTF-8."
        )

    if raw.startswith(b"%PDF") or name.endswith(".pdf"):
        return _normalize_newlines(_extract_pdf_text(raw))

    if zipfile.is_zipfile(io.BytesIO(raw)):
        if _is_xlsx_zip(raw) and not _is_docx_zip(raw):
            raise KeywordDatasetParseError("Excel .xlsx is not supported; use .docx or CSV.")
        if _is_docx_zip(raw) or name.endswith(".docx"):
            return _normalize_newlines(_docx_plain_text(raw))
        raise KeywordDatasetParseError("Unsupported ZIP upload. Use .docx or CSV.")

    return _normalize_newlines(_decode_bytes_for_csv(raw))


def ingest_article_keyword_dataset_bytes(raw: bytes, filename: str = "") -> list[dict[str, str]]:
    """
    Full admin pipeline: structured parse if possible; otherwise read full document text and
    call OpenAI to build {category, keyword} rows.
    """
    structured = _try_structured_keyword_rows(raw, filename)
    if structured and len(structured) > MAX_STRUCTURED_ROWS_BEFORE_AI:
        structured = []
    if structured:
        return structured[:MAX_DATASET_ROWS]

    text = _extract_full_document_text(raw, filename).strip()
    if len(text) < MIN_CHARS_FOR_AI_EXTRACTION:
        raise KeywordDatasetParseError(
            f"Not enough text to extract keywords automatically (need at least {MIN_CHARS_FOR_AI_EXTRACTION} characters). "
            "Try a longer PDF/DOCX, or upload CSV with columns category, keyword."
        )

    if len(text) > MAX_CHARS_SENT_TO_AI:
        text = text[:MAX_CHARS_SENT_TO_AI] + "\n\n[... truncated ...]"

    try:
        from api.services.openai_client import extract_membership_keywords_from_document

        rows = extract_membership_keywords_from_document(text, creative_seed=uuid.uuid4().hex[:14])
    except RuntimeError as e:
        msg = str(e)
        if "OPENAI_API_KEY" in msg:
            raise KeywordDatasetParseError(
                "Automatic keyword extraction needs OPENAI_API_KEY in the backend environment."
            ) from e
        raise KeywordDatasetParseError(msg) from e
    except ValueError as e:
        raise KeywordDatasetParseError(str(e)) from e

    if not rows:
        raise KeywordDatasetParseError("Automatic keyword extraction returned no usable rows.")
    return rows[:MAX_DATASET_ROWS]


def parse_keyword_dataset_bytes(raw: bytes, filename: str = "") -> list[dict[str, str]]:
    """Parse or ingest upload (structured + automatic AI extraction fallback)."""
    return ingest_article_keyword_dataset_bytes(raw, filename)


def parse_keyword_csv_file(path: str) -> list[dict[str, str]]:
    with open(path, "rb") as f:
        return ingest_article_keyword_dataset_bytes(f.read(), filename=path)


def _extract_pdf_text(raw: bytes) -> str:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=raw, filetype="pdf")
        try:
            parts = [doc.load_page(i).get_text() for i in range(doc.page_count)]
        finally:
            doc.close()
        return "\n".join(parts)
    except Exception:
        pass
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw), strict=False)
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception as exc:
        raise KeywordDatasetParseError(f"Could not read this PDF: {exc}") from exc


def _parse_docx_bytes(raw: bytes) -> list[dict[str, str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise KeywordDatasetParseError("python-docx is required for Word files.") from exc

    rows: list[dict[str, str]] = []
    try:
        doc = Document(io.BytesIO(raw))
    except Exception as exc:
        raise KeywordDatasetParseError(f"Could not read this Word file: {exc}") from exc

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                _append_row(rows, {"category": cells[0], "keyword": cells[1]})
            elif len(cells) == 1 and "," in cells[0]:
                a, b = cells[0].split(",", 1)
                _append_row(rows, {"category": a, "keyword": b})

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        low = line.lower()
        if low.startswith("category") and "keyword" in low:
            continue
        if "," in line:
            a, b = line.split(",", 1)
            _append_row(rows, {"category": a.strip(), "keyword": b.strip().strip('"')})

    return rows


def parse_keyword_csv_bytes(raw: bytes) -> list[dict[str, str]]:
    """
    Parse keyword CSV / plain text. Uses decoded text + newline='' pattern for csv.DictReader.
    """
    if not raw or not raw.strip():
        return []

    text = _decode_bytes_for_csv(raw)
    rows: list[dict[str, str]] = []
    dict_reader_kw: dict[str, Any] = {}
    if sys.version_info >= (3, 12):
        dict_reader_kw["strict"] = False

    stream = io.StringIO(text, newline="")
    try:
        reader = csv.DictReader(stream, **dict_reader_kw)
        if not reader.fieldnames:
            return []
        for row in reader:
            _append_row(rows, row)
        return rows
    except csv.Error:
        loose = _parse_keyword_csv_loose(text)
        return loose if loose else rows
    finally:
        stream.close()


def parse_keyword_dataset_bytes(raw: bytes, filename: str = "") -> list[dict[str, str]]:
    """
    Parse a keyword list from CSV/TXT, Word .docx, or PDF.
    Expects category + keyword (two columns, or lines like: business, some phrase).
    Legacy .doc (binary) is not supported.
    """
    if not raw or not raw.strip():
        return []

    name = (filename or "").lower()

    if raw.startswith(_OLE_DOC_MAGIC):
        raise KeywordDatasetParseError(
            "Legacy Word .doc files are not supported. In Word use File → Save As → "
            "Word Document (.docx) or CSV UTF-8, then upload again."
        )

    if raw.startswith(b"%PDF") or name.endswith(".pdf"):
        text = _extract_pdf_text(raw)
        rows = _parse_keyword_csv_loose(text)
        if not rows:
            raise KeywordDatasetParseError(
                "No keyword rows found in this PDF. Put one entry per line: category, keyword "
                "(for example: business, negotiation tactics)."
            )
        return rows

    if zipfile.is_zipfile(io.BytesIO(raw)):
        if _is_xlsx_zip(raw) and not _is_docx_zip(raw):
            raise KeywordDatasetParseError(
                "Excel .xlsx workbooks are not supported for this upload. "
                "Export as CSV UTF-8, or copy the table into Word (.docx) and upload that."
            )
        if _is_docx_zip(raw) or name.endswith(".docx"):
            rows = _parse_docx_bytes(raw)
            if not rows:
                raise KeywordDatasetParseError(
                    "No keyword rows found in this Word file. Use a table with columns "
                    "category and keyword, or one line per paragraph: business, example keyword."
                )
            return rows

    return parse_keyword_csv_bytes(raw)


def parse_keyword_csv_file(path: str) -> list[dict[str, str]]:
    with open(path, "rb") as f:
        return parse_keyword_dataset_bytes(f.read(), filename=path)


def dataset_category_counts(rows: list[dict[str, Any]] | None) -> dict[str, int]:
    out: dict[str, int] = {c: 0 for c in sorted(VALID_CATEGORIES)}
    if not rows:
        return out
    for r in rows:
        if not isinstance(r, dict):
            continue
        c = normalize_category(str(r.get("category") or ""))
        out[c] = out.get(c, 0) + 1
    return out
